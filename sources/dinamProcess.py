from sources.common.common import logger, processControl, log_

from sources.common.utils import huggingface_login
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
import torch
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer, util
from huggingface_hub import hf_hub_download
import time
import os

from llama_cpp import Llama

def getModel():
    filename="Meta-Llama-3-8B-Instruct-Q4_K_M.gguf"
    filePath = os.path.join(processControl.env['models'], filename)
    if os.path.exists(filePath):
        return filePath

    model_path = hf_hub_download(
        repo_id="bartowski/Meta-Llama-3-8B-Instruct-GGUF",
        filename=filePath,
        local_dir=processControl.env['models'],  # Custom directory
        local_dir_use_symlinks=False,  # Store actual file (not symlink)
        token=processControl.defaults['huggingface_login']  # Only needed if not logged in via CLI
    )
    return model_path

class EvaluadorLlama3Local:
    def __init__(self):
        # Load GGUF model (CPU version)
        self.model_path = "Meta-Llama-3-8B-Instruct-Q4_K_M.gguf"  # Download this file first
        self.model_path = getModel()
        self.llm = Llama(
            model_path=self.model_path,
            n_ctx=2048,  # Context window size
            n_threads=8,  # Use all CPU cores
            n_gpu_layers=0  # 0 = CPU-only mode
        )

        # Model for semantic similarity (unchanged)
        self.sim_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

    def extraer_texto(self, archivo):
        """Extrae texto de PDF o Word (unchanged)"""
        if archivo.endswith('.pdf'):
            with open(archivo, 'rb') as f:
                reader = PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages])
        elif archivo.endswith(('.docx', '.doc')):
            doc = Document(archivo)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Formato de archivo no soportado")

    def calcular_similitud(self, texto1, texto2):
        """Calcula similitud semántica entre textos (unchanged)"""
        emb1 = self.sim_model.encode(texto1, convert_to_tensor=True)
        emb2 = self.sim_model.encode(texto2, convert_to_tensor=True)
        return util.pytorch_cos_sim(emb1, emb2).item()

    def generar_respuesta(self, prompt):
        """Genera respuesta usando Llama 3 GGUF (modified for CPU)"""
        response = self.llm.create_chat_completion(
            messages=[{
                "role": "system",
                "content": "Eres un profesor asistente experto en evaluación académica de ciberseguridad."
            }, {
                "role": "user",
                "content": prompt
            }],
            max_tokens=512,
            temperature=0.7,
            top_p=0.9
        )
        return response['choices'][0]['message']['content']

    def evaluar_ejercicio(self, instrucciones, ejercicio):
        """Evalúa el ejercicio según las instrucciones (modified prompt format)"""

        summary_prompt = f"Resume este texto académico en menos de 1000 palabras:\n\n{ejercicio[:10000]}"  # First 10k chars

        summary = self.llm.create_chat_completion(
            messages=[{"role": "user", "content": summary_prompt}],
            max_tokens=1000
        )['choices'][0]['message']['content']

        ejercicio = summary
        toEliminate = "Aquí te presento un resumen del texto académico en menos de 1000 palabras:"
        similitud = self.calcular_similitud(instrucciones, ejercicio)

        prompt = f"""
        INSTRUCCIONES:
        {instrucciones}

        EJERCICIO DEL ESTUDIANTE:
        {ejercicio}

        Analiza este ejercicio y proporciona:
        1. Evaluación (Apto/No apto)
        2. Comentario con retroalimentación específica (máx. 50 palabras)
        3. Detección de posible uso de IA (Sí/No)
        """

        evaluacion = self.generar_respuesta(prompt)

        return {
            "evaluacion": evaluacion,
            "similitud": similitud
        }

    def procesar_lote(self, archivo_instrucciones, archivos_ejercicios):
        """Procesa múltiples ejercicios (unchanged)"""
        instrucciones = self.extraer_texto(archivo_instrucciones)
        resultados = []

        for archivo in archivos_ejercicios:
            try:
                ejercicio = self.extraer_texto(archivo)
                inicio = time.time()
                resultado = self.evaluar_ejercicio(instrucciones, ejercicio)
                tiempo = time.time() - inicio

                resultados.append({
                    "archivo": archivo,
                    "evaluacion": resultado["evaluacion"],
                    "similitud": resultado["similitud"],
                    "tiempo_procesamiento": tiempo
                })

            except Exception as e:
                print(f"Error procesando {archivo}: {str(e)}")
                resultados.append({
                    "archivo": archivo,
                    "error": str(e)
                })

        return resultados



class EvaluadorLlama3:
    def __init__(self):
        # Cargar modelo y tokenizer
        """
        self.model_name = "meta-llama/Meta-Llama-3-8B-Instruct-4bit"
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForCausalLM.from_pretrained(
            self.model_name,
            torch_dtype=torch.float16,
            device_map=processControl.defaults['device'],
            load_in_4bit=True  # ¡Importante para ahorrar memoria!
        )

        """
        self.model_name = "meta-llama/Meta-Llama-3-8B-Instruct"
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        # 4-bit quantization config
        bnb_config = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_use_double_quant=True,
            bnb_4bit_compute_dtype=torch.float16
        )
        self.model = AutoModelForCausalLM.from_pretrained(
            self.model_name,
            quantization_config=bnb_config,
            device_map=processControl.defaults['device'],
            torch_dtype=torch.float16
        )        


        # Modelo para similitud semántica
        self.sim_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

        # Configuración de generación
        self.gen_config = {
            "max_length": 2000,
            "temperature": 0.7,
            "top_p": 0.9,
            "do_sample": True
        }

    def extraer_texto(self, archivo):
        """Extrae texto de PDF o Word"""
        if archivo.endswith('.pdf'):
            with open(archivo, 'rb') as f:
                reader = PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages])
        elif archivo.endswith(('.docx', '.doc')):
            doc = Document(archivo)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Formato de archivo no soportado")

    def calcular_similitud(self, texto1, texto2):
        """Calcula similitud semántica entre textos"""
        emb1 = self.sim_model.encode(texto1, convert_to_tensor=True)
        emb2 = self.sim_model.encode(texto2, convert_to_tensor=True)
        return util.pytorch_cos_sim(emb1, emb2).item()

    def generar_respuesta(self, prompt):
        """Genera respuesta usando Llama 3 local"""
        inputs = self.tokenizer(prompt, return_tensors="pt").to(processControl.defaults['device'])

        with torch.no_grad():
            outputs = self.model.generate(
                **inputs,
                max_new_tokens=512,
                temperature=0.7,
                top_p=0.9,
                do_sample=True
            )

        return self.tokenizer.decode(outputs[0], skip_special_tokens=True)

    def evaluar_ejercicio(self, instrucciones, ejercicio):
        """Evalúa el ejercicio según las instrucciones"""
        # Primero calculamos similitud semántica
        similitud = self.calcular_similitud(instrucciones, ejercicio)

        # Construimos el prompt de evaluación
        prompt = f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>
        Eres un profesor asistente experto en evaluación académica. Analiza el ejercicio del estudiante comparándolo con las instrucciones proporcionadas.
        Advierte si detectas que se ha usado IA.

        INSTRUCCIONES:
        {instrucciones}

        EJERCICIO DEL ESTUDIANTE:
        {ejercicio}

        Proporciona:
        1. Evaluación (Apto/No apto) - basado en cumplimiento de requisitos
        2. Comentario final motivador, en caso de No apto identifica los fallos, en caso de Apto identifica sugerencias concretas de mejora(máx. 50 palabras)
        <|eot_id|><|start_header_id|>user<|end_header_id|>
        Por favor, provee la evaluación detallada según lo solicitado.<|eot_id|>"""

        # Generamos la evaluación
        evaluacion = self.generar_respuesta(prompt)

        return {
            "evaluacion": evaluacion,
            "similitud": similitud
        }

    def procesar_lote(self, archivo_instrucciones, archivos_ejercicios):
        """Procesa múltiples ejercicios"""
        a = 0
        instrucciones = self.extraer_texto(archivo_instrucciones)
        resultados = []

        for archivo in archivos_ejercicios:
            try:
                ejercicio = self.extraer_texto(archivo)
                inicio = time.time()

                resultado = self.evaluar_ejercicio(instrucciones, ejercicio)
                tiempo = time.time() - inicio

                resultados.append({
                    "archivo": archivo,
                    "evaluacion": resultado["evaluacion"],
                    "similitud": resultado["similitud"],
                    "tiempo_procesamiento": tiempo
                })

            except Exception as e:
                print(f"Error procesando {archivo}: {str(e)}")
                resultados.append({
                    "archivo": archivo,
                    "error": str(e)
                })

        return resultados


# Ejemplo de uso
def dinamizaProcess():
    processControl.defaults['device'] = "cuda" if torch.cuda.is_available() else "cpu"
    if processControl.defaults['device'] == "cuda":
        os.environ["RANK"] = "0"
        os.environ["WORLD_SIZE"] = "1"
        os.environ["MASTER_ADDR"] = "localhost"
        os.environ["MASTER_PORT"] = "12345"

        os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "expandable_segments:True"
        torch.cuda.set_per_process_memory_fraction(0.98, device=0)
        torch.backends.cuda.max_split_size_mb = 64
    else:
        os.environ["CUDA_VISIBLE_DEVICES"] = ""

    print("Inicializando modelo Llama 3... (esto puede tomar varios minutos)")

    instrucciones = os.path.join(processControl.env['inputPath'], processControl.args.act, f"{processControl.args.act}PlanteamientoYSolucionario.pdf")

    ejercicios = [
        os.path.join(processControl.env['inputPath'], processControl.args.act, "ejercicios", "parte1.pdf"),
        os.path.join(processControl.env['inputPath'], processControl.args.act, "ejercicios", "parte2.pdf")
        ]    

    #ejercicios = os.path.join(processControl.env['inputPath'], processControl.args.act, "ejercicios", "parte1.pdf")


    huggingface_login(processControl.defaults['huggingface_login'])
    evaluador = EvaluadorLlama3Local()
    log_("info", logger, "Modelo cargado correctamente")


    log_("info", logger, "\nProcesando ejercicios...")
    resultados = evaluador.procesar_lote(instrucciones, ejercicios)

    print("\nResultados de la evaluación:")
    for res in resultados:
        print(f"\nArchivo: {res['archivo']}")
        if 'error' in res:
            print(f"ERROR: {res['error']}")
        else:
            print(f"Similitud con instrucciones: {res['similitud']:.2f}")
            print(f"Tiempo procesamiento: {res['tiempo_procesamiento']:.2f}s")
            print("\nEVALUACIÓN:")
            log_("info", logger, res['evaluacion'])