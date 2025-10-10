from sources.common.common import logger, processControl, log_
import json

import time
from datetime import datetime
import os, re, unicodedata
from os.path import isdir
from PyPDF2 import PdfReader
from docx import Document
import shutil
from collections import Counter
import hashlib



def mkdir(dir_path):
    """
    @Desc: Creates directory if it doesn't exist.
    @Usage: Ensures a directory exists before proceeding with file operations.
    """
    if not isdir(dir_path):
        os.makedirs(dir_path)


def dbTimestamp():
    """
    @Desc: Generates a timestamp formatted as "YYYYMMDDHHMMSS".
    @Result: Formatted timestamp string.
    """
    timestamp = int(time.time())
    formatted_timestamp = str(time.strftime("%Y%m%d%H%M%S", time.gmtime(timestamp)))
    return formatted_timestamp

class configLoader:
    """
    @Desc: Loads and provides access to JSON configuration data.
    @Usage: Instantiates with path to config JSON file.
    """
    def __init__(self, config_path='config.json'):
        self.base_path = os.path.realpath(os.getcwd())
        realConfigPath = os.path.join(self.base_path, config_path)
        self.config = self.load_config(realConfigPath)

    def load_config(self, realConfigPath):
        with open(realConfigPath, 'r') as config_file:
            return json.load(config_file)

    def get_environment(self):
        environment =  self.config.get("environment", None)
        environment["realPath"] = self.base_path
        return environment

    def get_defaults(self):
        return self.config.get("defaults", {})

    def get_models(self):
        return self.config.get("models", {})

def image_parser(args):
    out = args.image_file.split(args.sep)
    return out

def huggingface_login(token):
    from huggingface_hub import login
    try:
        # Add your Hugging Face token here, or retrieve it from environment variables
        token = processControl.defaults['huggingface_login']
        login(token)
        print("Successfully logged in to Hugging Face.")
    except Exception as e:
        print("Error logging into Hugging Face:", str(e))
        raise


def extraer_texto(archivo):
    """Extrae texto de PDF o Word"""
    if archivo.endswith('.pdf'):
        with open(archivo, 'rb') as f:
            reader = PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages])
    elif archivo.endswith(('.docx', '.doc')):
        doc = Document(archivo)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Formato de archivo no soportado")


def limpiar_texto_para_llm(texto):
    """Limpia y normaliza texto para ser usado con LLMs."""
    texto = unicodedata.normalize("NFKC", texto)
    texto = re.sub(r"[\x00-\x1F\x7F]+", " ", texto)
    texto = re.sub(r"\s{2,}", " ", texto)

    lineas = texto.splitlines()
    limpio = []
    vistos = set()
    for linea in lineas:
        linea = linea.strip()
        if not linea or len(linea) < 3:
            continue
        if linea in vistos:
            continue
        vistos.add(linea)
        limpio.append(linea)
    texto = "\n".join(limpio)
    texto = re.sub(r"\n{2,}", "\n", texto).strip()
    return texto


def extraerTextoMetas(archivo, max_paginas=None):
    """
    Extrae texto limpio y metadatos básicos, preparado para enviar a un LLM.
    """
    resultado = {"texto": "", "metadatos": {}}
    nombre = os.path.basename(archivo)
    extension = os.path.splitext(nombre)[1].lower()

    resultado["metadatos"].update({
        "nombre_archivo": nombre,
        "tamano_bytes": os.path.getsize(archivo),
        "extension": extension,
    })

    texto_completo = ""

    if extension == ".pdf":
        try:
            with open(archivo, "rb") as f:
                reader = PdfReader(f)
                num_paginas = len(reader.pages)
                resultado["metadatos"]["num_paginas"] = num_paginas

                for i, page in enumerate(reader.pages):
                    if max_paginas and i >= max_paginas:
                        break
                    try:
                        page_text = page.extract_text() or ""
                        texto_completo += f"\n--- Página {i+1} ---\n{page_text}"
                    except Exception as e:
                        texto_completo += f"\n[Error en página {i+1}: {e}]"
        except Exception as e:
            texto_completo = f"[ERROR leyendo PDF: {e}]"

    elif extension == ".docx":
        try:
            doc = Document(archivo)
            texto = "\n".join(para.text for para in doc.paragraphs)
            texto_completo = texto

            props = doc.core_properties
            resultado["metadatos"].update({
                "titulo": props.title or "",
                "autor": props.author or "",
                "creado": str(props.created or ""),
                "ultima_modificacion": str(props.modified or "")
            })
        except Exception as e:
            texto_completo = f"[ERROR leyendo DOCX: {e}]"

    else:
        raise ValueError("Solo se admiten .pdf o .docx")

    # Limpieza final
    texto_limpio = limpiar_texto_para_llm(texto_completo)
    resultado["texto"] = texto_limpio
    return resultado


def grabaJson(data, path):
    try:

        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

    except Exception as e:
        log_("error", logger, f"Error write json path:{path}, error:{e}")
        return False

    log_("info", logger, f"JSON written path:{path}")
    return True


def clean_and_move(path, filepath1, filepath2):
    # Validate inputs
    if not os.path.exists(path):
        print(f"Error: Path {path} does not exist")
        return False

    try:
        # Recursively delete all contents under path
        for root, dirs, files in os.walk(path, topdown=False):
            for name in files:
                file_path = os.path.join(root, name)
                try:
                    os.unlink(file_path)
                except Exception as e:
                    raise Exception(f"Failed to delete {file_path}: {e}")


            for name in dirs:
                dir_path = os.path.join(root, name)
                try:
                    os.rmdir(dir_path)
                except Exception as e:
                    raise Exception(f"Failed to delete {dir_path}: {e}")

        print(f"Successfully cleaned directory: {path}")

        if filepath1 is None:
            return True

        if not os.path.exists(filepath1):
            raise Exception(f"Error: Source file {filepath1} does not exist")

        if not os.path.exists(filepath2):
            raise Exception(f"Error: Source file {filepath2} does not exist")

        # Move filepath1 to filepath2
        shutil.move(filepath1, filepath2)
        print(f"Successfully moved {filepath1} to {filepath2}")

        return True

    except Exception as e:
        log_("exception", logger, f"Operation failed: {e}")
        return False

def determinarTema(texto, actividad):
    # Palabras clave para cada tema (puedes expandirlas según necesidad)
    """
    identificador = {
        "actividad1":{
            "ciberinteligencia":["ciberinteligencia", "inteligencia", "osint", "amenazas", "estrategica", "tactica", "soc", "vigilancia"],
            "ransomware":["ransomware", "malware", "cifrado", "rescate", "eternalblue", "wannacry", "bitcoin", "secuestro"]
        },
        "actividad2":{
            "capas":["capas", "seguridad por capas"],
            "teletrabajo":["teletrabajo"]
        }
    }    
    """

    identificador = processControl.defaults['identificador']
    texto = texto.lower()

    # Crear un contador para cada subtema
    pesos = {}

    # Iterar sobre los subtemas del nivel actual
    for subtema, palabras_clave in identificador.get(actividad, {}).items():
        peso = 0
        for palabra in palabras_clave:
            # Contar ocurrencias exactas o frases (como "seguridad por capas")
            if " " in palabra:
                # Frase exacta
                peso += texto.count(palabra.lower())
            else:
                # Palabra suelta (usamos expresiones regulares para evitar coincidencias parciales)
                ocurrencias = re.findall(rf'\b{re.escape(palabra.lower())}\b', texto)
                peso += len(ocurrencias)
        pesos[subtema] = peso

    # Si todos los pesos son 0, no se puede determinar el tema
    if all(p == 0 for p in pesos.values()):
        return "indeterminado"

    # Devolver el subtema con mayor peso
    return max(pesos, key=pesos.get)


    """
    palabras_ciberinteligencia = {"ciberinteligencia", "inteligencia", "osint", "amenazas", "estrategica", "tactica", "soc", "vigilancia"}
    palabras_ransomware = {"ransomware", "malware", "cifrado", "rescate", "eternalblue", "wannacry", "bitcoin", "secuestro"}

    # Convertir texto a minúsculas y eliminar puntuación
    texto_limpio = re.sub(r'[^\w\s]', '', texto.lower())
    palabras_texto = texto_limpio.split()

    # Contar coincidencias
    contador_ci = Counter(palabras_texto) & Counter(palabras_ciberinteligencia)
    contador_r = Counter(palabras_texto) & Counter(palabras_ransomware)

    # Calcular puntuación (similitud basada en número de palabras clave)
    similitud_ci = sum(contador_ci.values()) / len(palabras_ciberinteligencia) if palabras_ciberinteligencia else 0
    similitud_r = sum(contador_r.values()) / len(palabras_ransomware) if palabras_ransomware else 0

    # Determinar tema con mayor similitud
    if similitud_ci > similitud_r and similitud_ci > 0.1:  # Umbral mínimo de 10% para certeza
        return "ciberinteligencia", similitud_ci
    elif similitud_r > similitud_ci and similitud_r > 0.1:
        return "ransomware", similitud_r
    else:
        return "desconocido", max(similitud_ci, similitud_r)    
    """
def extraer_json_de_texto(output_text):
    """
    Extrae el primer bloque JSON válido desde el texto completo generado por el modelo.
    """
    posibles_jsones = re.findall(r'\{.*?\}', output_text, re.DOTALL)
    for j in posibles_jsones:
        try:
            return json.loads(j)
        except json.JSONDecodeError:
            continue
    return None


def read_json(filepath):
    """
    Reads a JSON file and returns its contents.
    If the file does not exist or is empty, returns [].
    """
    if not os.path.exists(filepath):
        return []

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data
    except (json.JSONDecodeError, ValueError):
        # File exists but is not valid JSON
        return []


def searchAlumnoHistorico(json_data, alumno_name):
    """
    Search in a JSON list of dicts for entries where 'alumno' == alumno_name.

    Args:
        json_data (list): List of dictionaries with keys "alumno" and "ejercicio".
        alumno_name (str): Name of the alumno to search for.

    Returns:
        list: All matching elements. Empty list if none found.
    """
    if not isinstance(json_data, list):
        return []

    results = [item for item in json_data if item.get("alumno") == alumno_name]
    return results if results else []


def upsertAlumnoHistorico(json_data, elemento):
    """
    Inserta o actualiza el registro de un alumno en el histórico.

    Si el alumno ya existe en json_data, se eliminan sus entradas anteriores
    y se inserta la nueva. Si no existe, simplemente se inserta.

    Args:
        json_data (list): Lista de diccionarios con claves 'alumno' y 'ejercicio'.
        alumno_name (str): Nombre del alumno.
        ejercicio (int): Número de ejercicio.

    Returns:
        list: La lista actualizada.
    """
    if not isinstance(json_data, list):
        json_data = []

    # Eliminar entradas previas del alumno
    json_data = [item for item in json_data if item.get("alumno") != elemento['alumno']]

    # Insertar el nuevo registro
    json_data.append(elemento)

    return json_data


def texto_hash(texto, method="sha256"):
    """
    Devuelve el hash de un texto.
    """
    if method == "md5":
        return hashlib.md5(texto.encode("utf-8")).hexdigest()
    elif method == "sha256":
        return hashlib.sha256(texto.encode("utf-8")).hexdigest()
    else:
        raise ValueError("Método de hash no soportado")


def existe_hash_en_otro_alumno(json_data, alumnoBuscar, hashBuscar):
    """
    Busca si un hash ya existe en el JSON en registros de alumnos distintos al alumno dado.

    Args:
        json_data (list): Lista de alumnos con sus resultados.
        alumnoBuscar (str): Nombre del alumno actual.
        hashBuscar (str): Hash a verificar.

    Returns:
        dict or None: El registro encontrado (alumno, tema, archivo) si existe en otro alumno,
                      None si no hay coincidencias.
    """
    for alumno_entry in json_data:
        alumno = alumno_entry.get("alumno")
        if alumno != alumnoBuscar:  # mirar solo en otros alumnos
            for resultado in alumno_entry.get("resultados", []):
                if resultado.get("hash") == hashBuscar:
                    return {
                        "alumno": alumno,
                        "tema": resultado.get("tema"),
                        "archivo": resultado.get("archivo")
                    }
    return None


def calcular_semaforos(metasDict):
    """
    Genera un diccionario de 'semaforos' a partir de metadatos.

    Args:
        metasDict (dict): Diccionario con posibles claves como 'autor', 'Author',
                          'ultima_modificacion', 'creado', 'Producer', 'Creator'.

    Returns:
        dict: Diccionario con los semáforos construidos.
    """
    semaforos = {}

    # Autor (case-insensitive entre autor / Author)
    if "autor" in metasDict and metasDict["autor"]:
        semaforos["autor"] = metasDict["autor"]
    elif "Author" in metasDict and metasDict["Author"]:
        semaforos["autor"] = metasDict["Author"]

    # Diferencia temporal en minutos (truncada a 2 decimales)
    if "ultima_modificacion" in metasDict and "creado" in metasDict:
        try:
            fmt = "%Y-%m-%d %H:%M:%S%z"  # formato "2025-07-04 10:26:00+00:00"
            mod_time = datetime.strptime(str(metasDict["ultima_modificacion"]), fmt)
            creado_time = datetime.strptime(str(metasDict["creado"]), fmt)
            diff_minutes = (mod_time - creado_time).total_seconds() / 60
            semaforos["time"] = round(diff_minutes, 2)  # truncado a 2 decimales
        except Exception as e:
            semaforos["time"] = f"Error calculando diferencia: {e}"

    # Producer o Creator (si existen, ambos van a 'producer')
    if "Producer" in metasDict and metasDict["Producer"]:
        semaforos["producer"] = metasDict["Producer"]
    elif "Creator" in metasDict and metasDict["Creator"]:
        semaforos["producer"] = metasDict["Creator"]

    return semaforos